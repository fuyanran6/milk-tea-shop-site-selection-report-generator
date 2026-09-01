"""Word DOCX export with embedded analysis map."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT_NAME = "Microsoft YaHei"
META_GRAY = RGBColor(0x86, 0x8E, 0x96)


def export_docx(report: dict[str, Any], png_path: Path, output_path: Path, meta: dict[str, Any]) -> Path:
    doc = Document()
    for style_name in ("Normal", "Heading 1", "Heading 2", "List Bullet", "List Number"):
        if style_name in doc.styles:
            _set_style_font(doc.styles[style_name])

    title = doc.add_heading("奶茶店选址可行性商业分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.bold = True
        _style_run(run, size=22)

    p1 = doc.add_paragraph(f"候选地址：{meta.get('address', '')}")
    _style_paragraph(p1)
    if meta.get("is_demo"):
        demo_p = doc.add_paragraph("【演示点报告】本报告基于内置演示数据生成，用于产品功能展示与体验，不代表对真实候选地址的分析结论。")
        _style_paragraph(demo_p)
        for run in demo_p.runs:
            run.bold = True
    p2 = doc.add_paragraph(f"生成时间：{meta.get('generated_at', '')}")
    _style_paragraph(p2)

    for ch in report.get("chapter_order", []):
        key = ch["key"]
        chapter = report["chapters"].get(key, {})
        h = doc.add_heading(chapter.get("title", ch["title"]), level=1)
        _style_heading(h)
        headline = chapter.get("headline")
        if headline:
            p = doc.add_paragraph()
            run = p.add_run(headline)
            run.bold = True
            _style_run(run, size=11)
        _write_content(doc, chapter.get("content", ""))

        if key == "map" and png_path.exists():
            doc.add_picture(str(png_path), width=Inches(5.5))

    h = doc.add_heading("证据清单（摘要）", level=1)
    _style_heading(h)
    for ev in report.get("evidence", [])[:20]:
        p = doc.add_paragraph(f"{ev['claim_id']}：{ev['claim']}（来源：{ev['source']}）")
        _style_paragraph(p)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def _set_style_font(style, size=11):
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)


def _style_run(run, size=11, gray=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT_NAME)
    if gray:
        run.font.color.rgb = META_GRAY


def _style_paragraph(paragraph, gray=False):
    for run in paragraph.runs:
        _style_run(run, gray=gray)


def _style_heading(paragraph):
    for run in paragraph.runs:
        _style_run(run, size=14)


def _is_meta_gray_line(line: str) -> bool:
    if not line:
        return False
    if "图层说明" in line:
        return True
    if "数据来源" in line:
        return True
    if line.startswith("街道底图：") or line.startswith("未提供高德") or line.startswith("高德静态"):
        return True
    low = line.lower()
    if "osm" in low and any(k in line for k in ("缓存", "覆盖", "暂未获取", "稀疏", "建筑轮廓")):
        return True
    if line.startswith("<p class=\"kb-ref\">") or line.startswith("<p class='kb-ref'>"):
        return True
    return False


def _add_text_paragraph(doc: Document, text: str, *, bullet=False, number=False, gray=False):
    style = None
    if bullet:
        style = "List Bullet"
    elif number:
        style = "List Number"
    p = doc.add_paragraph(text, style=style)
    if not p.runs:
        p.add_run(text)
    _style_paragraph(p, gray=gray)
    return p


def _write_content(doc: Document, content: str) -> None:
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for i, row in enumerate(table_rows):
            for j in range(cols):
                text = row[j] if j < len(row) else ""
                table.rows[i].cells[j].text = text
                for p in table.rows[i].cells[j].paragraphs:
                    _style_paragraph(p)
        table_rows = []

    for raw in content.split("\n"):
        line = raw.strip()
        if not line:
            flush_table()
            continue
        if line.startswith("<p class=\"kb-ref\">") or line.startswith("<p class='kb-ref'>"):
            line = re.sub(r"</?p[^>]*>", "", line)
            _add_text_paragraph(doc, line, gray=True)
            continue
        if line.startswith("<") and line.endswith(">"):
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(set(c) <= set("-: ") and c for c in cells):
                continue
            table_rows.append(cells)
            continue

        flush_table()
        gray = _is_meta_gray_line(line)

        if line.startswith("### "):
            h = doc.add_heading(_strip_md(line[4:]), level=2)
            _style_heading(h)
        elif line.startswith("## "):
            h = doc.add_heading(_strip_md(line[3:]), level=2)
            _style_heading(h)
        elif line.startswith("- "):
            _add_text_paragraph(doc, _strip_md(line[2:]), bullet=True, gray=gray)
        elif re.match(r"^\d+\.\s+", line):
            _add_text_paragraph(doc, _strip_md(re.sub(r"^\d+\.\s+", "", line)), number=True, gray=gray)
        else:
            _add_text_paragraph(doc, _strip_md(line), gray=gray)

    flush_table()


def _strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = text.replace("*", "")
    return text
